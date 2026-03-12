from pathlib import Path


def parse_resume_file(file_path: str) -> dict:
    path = Path(file_path)
    if not path.exists():
        return {"success": False, "error": f"File not found: {file_path}", "content": ""}

    ext = path.suffix.lower()

    try:
        if ext == ".txt":
            text = path.read_text(encoding="utf-8", errors="replace")

        elif ext == ".docx":
            from docx import Document
            doc = Document(str(path))

            parts = []
            for para in doc.paragraphs:
                t = (para.text or "").strip()
                if t:
                    parts.append(t)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = (cell.text or "").strip()
                        if t:
                            parts.append(t)

            text = "\n".join(parts)

        elif ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(path))

            parts = []
            for page in reader.pages:
                t = (page.extract_text() or "").strip()
                if t:
                    parts.append(t)

            text = "\n".join(parts)

        else:
            return {"success": False, "error": f"Unsupported format '{ext}'", "content": ""}

        text = (text or "").strip()
        if len(text.split()) < 50:
            return {
                "success": False,
                "error": "Parsed content too short — file may be empty or image-only.",
                "content": text,
            }

        return {
            "success": True,
            "content": text,
            "file_name": path.name,
            "file_type": ext.lstrip("."),
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Parse error: {type(e).__name__}: {e}",
            "content": "",
        }
